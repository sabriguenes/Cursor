"""
Generate McKinsey-style Word document for the Q&A Architecture Decisions.
Run: python generate_qa_docx.py
Output: QA-ARCHITECTURE-DECISIONS.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Color Palette (same as Consultant Paper) ──
NAVY = RGBColor(0x00, 0x33, 0x66)
DARK_BLUE = RGBColor(0x00, 0x52, 0x8A)
ELECTRIC = RGBColor(0x00, 0x96, 0xD6)
CHARCOAL = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
WARM_GRAY = RGBColor(0x6B, 0x6B, 0x6B)
GREEN_ACCENT = RGBColor(0x00, 0x7A, 0x33)
RED_ACCENT = RGBColor(0xCC, 0x00, 0x00)

TBL_HEADER_BG = "003366"
TBL_ALT_BG = "F0F4F8"
TBL_BORDER_COLOR = "B0BEC5"
ACCENT_BAR_COLOR = "0096D6"


def set_keep_with_next(paragraph):
    """Keep this paragraph on the same page as the next one."""
    pPr = paragraph._p.get_or_add_pPr()
    keep = parse_xml(f'<w:keepNext {nsdecls("w")} />')
    pPr.append(keep)


def set_keep_together(paragraph):
    """Prevent this paragraph from splitting across pages."""
    pPr = paragraph._p.get_or_add_pPr()
    keep = parse_xml(f'<w:keepLines {nsdecls("w")} />')
    pPr.append(keep)


def set_page_break_before(paragraph):
    """Start this paragraph on a new page WITHOUT adding an empty paragraph.
    Unlike doc.add_page_break(), this never creates blank pages."""
    pPr = paragraph._p.get_or_add_pPr()
    pb = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} />')
    pPr.append(pb)


def set_cell_bg(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="B0BEC5", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(cell, TBL_HEADER_BG)
        set_cell_borders(cell, TBL_HEADER_BG, "6")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(val)
            if text.startswith("**") and text.endswith("**"):
                run = p.add_run(text.strip("*"))
                run.bold = True
            else:
                run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = CHARCOAL
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_borders(cell, TBL_BORDER_COLOR)
            if r_idx % 2 == 1:
                set_cell_bg(cell, TBL_ALT_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def add_accent_bar(doc, sz="12"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{ACCENT_BAR_COLOR}"/>'
        f'</w:pBorders>'
    )
    pPr.append(borders)
    p.space_after = Pt(6)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = NAVY
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = ELECTRIC
            run.font.size = Pt(13)
        run.font.name = "Calibri"
    return h


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = color or CHARCOAL
    run.bold = bold
    run.italic = italic
    p.space_after = Pt(4)
    p.space_before = Pt(2)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{ACCENT_BAR_COLOR}"/>'
        f'</w:pBorders>'
    )
    pPr.append(borders)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720"/>')
    pPr.append(ind)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = WARM_GRAY
    run.italic = True
    p.space_after = Pt(8)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shading)
    borders = parse_xml(
        f'<w:pBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="DDDDDD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="DDDDDD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'</w:pBorders>'
    )
    pPr.append(borders)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = "Consolas"
    run.font.color.rgb = CHARCOAL
    p.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = CHARCOAL
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = CHARCOAL
    return p


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph().space_after = Pt(20)
    add_accent_bar(doc, sz="48")  # THICK bar — visible on cover

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Architecture Decisions")
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    run.bold = True
    p.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Questions & Answers\nMulti-Agent Code Review System for LLVM/Clang")
    run.font.size = Pt(20)
    run.font.color.rgb = ELECTRIC
    run.font.name = "Calibri"
    p.space_after = Pt(30)

    add_accent_bar(doc, sz="48")  # THICK bar — visible on cover

    meta = [
        ("Document Type", "Internal Q&A Reference"),
        ("Date", "February 10, 2026"),
        ("Context", "Multi-Agent Code Review System for LLVM/Clang"),
        ("Version", "1.0"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}:  ")
        run.font.size = Pt(10)
        run.font.color.rgb = WARM_GRAY
        run.font.name = "Calibri"
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = CHARCOAL
        run.font.name = "Calibri"
        p.space_after = Pt(2)
    doc.add_page_break()


def build_q1(doc):
    add_heading_styled(doc, "Q1: Why Claude Code and not Cursor CLI?", 1)
    add_accent_bar(doc)

    add_body(doc, "Claude Code has a native Python SDK, built-in subagents, and 15 hook events. Cursor CLI has none of these.", bold=True)

    add_body(doc, (
        "Important correction: Cursor CLI does not perform codebase indexing -- only the Cursor IDE does "
        "(source: official Cursor docs, docs.cursor.com/context/codebase-indexing, confirmed by Cursor team "
        "member Tee on Discord). Both CLIs work file-based (Read, Grep, Glob) without an index. The 2-5 minute "
        "delay in the current setup comes from GitHub Runner allocation + npm install, not from any CLI indexing."
    ), italic=True, color=RED_ACCENT)

    add_heading_styled(doc, "The Actual Differentiators", 3)

    add_styled_table(doc,
        ["Capability", "Cursor CLI (cursor-agent)", "Claude Code / Agent SDK", "Winner"],
        [
            ["Codebase indexing", "No (IDE-only feature)", "No", "Tie"],
            ["How it finds code", "File-based (Read, Grep, Glob)", "File-based (Read, Grep, Glob)", "Tie"],
            ["Session persistence", "--resume (Beta)", "--resume + --continue", "Claude Code"],
            ["MCP support", "Yes (via mcp.json)", "Yes (native)", "Tie"],
            ["Subagents / Agent Teams", "Not available natively", "Built-in feature", "Claude Code"],
            ["Python SDK", "Not available", "claude-agent-sdk (pip install)", "Claude Code"],
            ["Hook system", "8 events (IDE-only)", "15 events (CLI-compatible)", "Claude Code"],
            ["Daemon mode", "No", "No -- but SDK embeds into FastAPI", "Claude Code"],
            ["Multi-model support", "GPT-5, Claude, Gemini, Grok", "Anthropic-only", "Cursor CLI"],
        ],
        col_widths=[3.5, 4.5, 4.5, 3.5]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_body(doc, "The decision is based on three capabilities Cursor CLI lacks entirely:", bold=True)
    add_bullet(doc, " -- agents defined as Python objects, callable natively from FastAPI", "Python SDK")
    add_bullet(doc, " -- built-in feature for multi-agent orchestration without external frameworks", "Subagents")
    add_bullet(doc, " -- PreToolUse, PostToolUse, SessionStart, Stop, etc., all functional in headless mode", "15 Hook Events")

    doc.add_paragraph().space_after = Pt(4)
    add_body(doc, (
        "Cursor CLI is an excellent tool for interactive development, but it lacks the programmatic "
        "orchestration layer we need for a webhook-driven multi-agent system."
    ))

    h = add_heading_styled(doc, "Sources", 3)
    set_keep_with_next(h)
    b1 = add_bullet(doc, 'docs.cursor.com/context/codebase-indexing -- "Cursor indexes your codebase... When you open a project" (IDE feature)')
    set_keep_with_next(b1)
    b2 = add_bullet(doc, "docs.cursor.com/en/cli/overview -- no mention of indexing")
    set_keep_with_next(b2)
    b3 = add_bullet(doc, 'Discord: Tee (Cursor team) -- "Cursor CLI does not index codebase, this is only done by the IDE"')
    set_keep_with_next(b3)
    b4 = add_bullet(doc, "code.claude.com/docs/en/headless -- programmatic usage with SDK")
    set_keep_with_next(b4)
    add_bullet(doc, "code.claude.com/docs/en/hooks-guide -- 15 event types")


def build_q2(doc):
    h = add_heading_styled(doc, "Q2: Can every developer use the webhook server?", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "Yes. Any developer with push access to the GitHub repo automatically triggers the webhook server. No local Claude Code installation required.", bold=True)

    add_heading_styled(doc, "How It Works", 3)

    flow = (
        "Developer pushes code / creates PR / comments on issue\n"
        "    |\n"
        "    v\n"
        "GitHub fires a webhook (HTTP POST) to our FastAPI server\n"
        "    |\n"
        "    v\n"
        "FastAPI server receives the event, validates the signature\n"
        "    |\n"
        "    v\n"
        "Server spawns the appropriate agent(s) via Claude Agent SDK\n"
        "    |\n"
        "    v\n"
        "Agent works, posts result back to GitHub as a comment/review"
    )
    add_code_block(doc, flow)

    add_heading_styled(doc, "Key Points", 3)
    add_bullet(doc, " -- developers interact only with GitHub, the server handles everything", "No local setup required")
    add_bullet(doc, " -- managed centrally on the server, not per-developer", "Single ANTHROPIC_API_KEY")
    add_bullet(doc, " -- visible in the monitoring dashboard", "Token costs tracked per agent run")
    add_bullet(doc, " -- dashboard can expose a manual trigger endpoint", "Optional direct access")
    add_bullet(doc, " -- GitHub webhook signatures ensure only legitimate events are processed", "Access control")


def build_q3(doc):
    h = add_heading_styled(doc, "Q3: Does it run 24/7?", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "The server runs 24/7. The agents run on-demand.", bold=True)

    add_heading_styled(doc, "What Runs Permanently (24/7)", 3)

    add_styled_table(doc,
        ["Component", "Runtime", "Monthly Cost"],
        [
            ["FastAPI webhook server", "24/7 (systemd service)", "$0 (on Azure VM)"],
            ["MCP server (Pinecone)", "24/7 (systemd service)", "$0 (on Azure VM)"],
            ["SQLite database", "24/7 (file on disk)", "$0"],
            ["nginx reverse proxy", "24/7 (HTTPS termination)", "$0"],
            ["Azure VM (B2ms)", "24/7", "~$60/month (MS credits)"],
        ],
        col_widths=[5, 5, 5]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_heading_styled(doc, "What Runs On-Demand (Event-Triggered)", 3)

    add_styled_table(doc,
        ["Component", "Trigger", "Duration", "Cost"],
        [
            ["Knowledge Agent", "GitHub event arrives", "Seconds", "~$0.01-0.05"],
            ["Coding Agent", "Knowledge agent returns context", "Minutes", "~$0.50-2.00"],
            ["Review Agent", "Coding agent completes changes", "Seconds-minutes", "~$0.10-0.30"],
        ],
        col_widths=[4, 4, 3.5, 3.5]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_quote(doc, (
        "Analogy: A doctor on call -- not operating 24/7, but reachable 24/7. "
        "The server is the hospital that never closes. The agents are the specialists called in when needed."
    ))

    add_heading_styled(doc, "Reliability", 3)
    add_bullet(doc, "systemd ensures auto-restart if the FastAPI server crashes")
    add_bullet(doc, "Health check endpoint (/health) for external monitoring")
    add_bullet(doc, "Azure VM uptime SLA: 99.9%")
    add_bullet(doc, "GitHub webhooks queue and retry automatically on failure")


def build_q4(doc):
    h = add_heading_styled(doc, "Q4: Is it cost-efficient?", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "Yes, through three cost levers: hybrid LLM strategy, on-demand spawning, and Microsoft startup credits.", bold=True)

    add_heading_styled(doc, "Lever 1: Hybrid LLM Strategy", 3)

    add_styled_table(doc,
        ["Agent", "Model", "Why This Model", "Cost per Invocation"],
        [
            ["Knowledge Agent", "Sonnet 4", "Simple retrieval, no code generation", "~$0.01-0.05 (cheapest)"],
            ["Coding Agent", "Opus 4.6", "C++ quality demands the best model", "~$0.50-2.00 (most expensive)"],
            ["Review Agent", "Sonnet 4", "Review quality, fewer tokens", "~$0.10-0.30 (medium)"],
        ],
        col_widths=[3, 2.5, 5.5, 5]
    )

    doc.add_paragraph().space_after = Pt(4)
    add_body(doc, (
        "The cost gradient works naturally: the knowledge agent runs most frequently (every event) "
        "but is the cheapest. The expensive coding agent only runs when actual code changes are needed."
    ))

    add_heading_styled(doc, "Lever 2: On-Demand Spawning", 3)
    add_bullet(doc, "Agents are spawned per-event, not running continuously")
    add_bullet(doc, "Zero API cost when no GitHub events arrive")
    add_bullet(doc, "No wasted tokens on idle processes")
    add_bullet(doc, "max_budget_usd parameter prevents token overruns per invocation")

    add_heading_styled(doc, "Lever 3: Microsoft Startup Credits ($5,000)", 3)

    add_styled_table(doc,
        ["Component", "Monthly Cost", "Covered By", "Runway"],
        [
            ["Azure VM (B2ms)", "~$60", "MS credits", "83 months"],
            ["Azure Blob Storage", "~$5", "MS credits", "--"],
            ["Azure Key Vault", "~$1", "MS credits", "--"],
            ["Total infrastructure", "~$66/month", "MS credits", "76+ months"],
        ],
        col_widths=[4, 3, 4, 4]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_heading_styled(doc, "Comparison with Alternatives", 3)

    add_styled_table(doc,
        ["Approach", "Monthly Cost", "Why More/Less Expensive"],
        [
            ["Our architecture", "~$66 + API tokens", "Minimal infra, on-demand agents"],
            ["GitHub Runners (current)", "$0 infra + API + dev time", "Free infra but 2-5 min/interaction = costly in dev hours"],
            ["LangGraph/CrewAI stack", "~$200+ + API tokens", "Additional SaaS fees, complexity overhead"],
        ],
        col_widths=[4, 4, 8]
    )


def build_q5(doc):
    h = add_heading_styled(doc, "Q5: Is a token/prompt strategy defined?", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "Yes. Every agent has scoped prompts, restricted tools, turn limits, and budget caps.", bold=True)

    add_heading_styled(doc, "Layer 1: Scoped System Prompts", 3)
    add_body(doc, "Each agent receives ONLY the instructions relevant to its role. No agent sees the full system context:")
    add_bullet(doc, ' "Query Pinecone and SQLite. Return structured context. DO NOT write code."', "Knowledge Agent: ")
    add_bullet(doc, ' "Implement changes based on context. Clean diffs. Clear commits."', "Coding Agent: ")
    add_bullet(doc, ' "Review changes. Respond ALL CLEAR or return specific feedback."', "Review Agent: ")

    add_heading_styled(doc, "Layer 2: Tool Restrictions", 3)

    add_styled_table(doc,
        ["Agent", "Allowed Tools", "Forbidden"],
        [
            ["Knowledge Agent", "Read, Grep, Glob, MCP", "Edit, Write, Bash, Git"],
            ["Coding Agent", "Read, Edit, Write, Bash, Grep, Glob", "-- (full access)"],
            ["Review Agent", "Read, Grep, Glob, Bash (tests only)", "Edit, Write"],
            ["Orchestrator", "None (delegates only)", "ALL coding tools"],
        ],
        col_widths=[3.5, 6, 5.5]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_heading_styled(doc, "Layer 3: Execution Limits", 3)

    add_styled_table(doc,
        ["Parameter", "Value", "Purpose"],
        [
            ["max_turns per agent", "10-20 (configurable)", "Prevent infinite loops"],
            ["max_budget_usd per agent", "$2.00 (configurable)", "Hard cost ceiling per invocation"],
            ["Review loop iterations", "Max 3", "Prevent endless coding/review cycles"],
            ["Context window", "Each agent starts fresh", "No accumulated bloat across agents"],
        ],
        col_widths=[4, 4, 8]
    )

    doc.add_paragraph().space_after = Pt(6)
    add_heading_styled(doc, "Layer 4: Context Passing (Not Sharing)", 3)

    flow = (
        "Knowledge Agent -> returns: structured context summary (500-1000 tokens)\n"
        "    |\n"
        "    v (only the summary is passed, not the full retrieval)\n"
        "Coding Agent -> returns: list of changed files + diff (variable)\n"
        "    |\n"
        "    v (only the diff is passed, not the full codebase context)\n"
        'Review Agent -> returns: "ALL CLEAR" or specific feedback (100-500 tokens)'
    )
    add_code_block(doc, flow)

    add_body(doc, (
        "Each agent operates within its own context window, never inheriting bloated context from "
        "a previous agent. Token efficiency is maximized by design."
    ))


def build_q6(doc):
    h = add_heading_styled(doc, "Q6: Are there professionally specialized agent roles?", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "Yes. 3 specialized agents + 1 orchestrator, each with strict role boundaries, scoped tools, and dedicated model selection.", bold=True)

    add_heading_styled(doc, "The Four Roles", 3)

    add_styled_table(doc,
        ["Role", "Specialization", "Tools", "Model", "Cost Tier"],
        [
            ["Orchestrator", "Delegates ONLY, codes NEVER", "No coding tools", "FastAPI logic", "$0 (no LLM)"],
            ["Knowledge Agent", "C++ knowledge retrieval", "Read, Grep, Glob, MCP", "Sonnet 4", "LOW"],
            ["Coding Agent", "Senior C++ developer", "Read, Edit, Write, Bash, Git", "Opus 4.6", "HIGH"],
            ["Review Agent", "Code reviewer / QA", "Read, Grep, Glob, Bash", "Sonnet 4", "MEDIUM"],
        ],
        col_widths=[3, 3.5, 3.5, 3, 3]
    )

    doc.add_paragraph().space_after = Pt(6)

    add_quote(doc, (
        '"CRITICAL: YOU ARE FORBIDDEN FROM CALLING AGENTS OTHER THAN THE ONES LISTED. '
        'Your ONLY job is to delegate to these agents. THAT\'S IT." '
        '-- Adapted from SteffenBlake/Atomic.Net manager.agent.md'
    ))

    add_heading_styled(doc, "Pipeline Flow", 3)

    pipeline = (
        "GitHub Event\n"
        "    |\n"
        "    v\n"
        '[1] KNOWLEDGE AGENT  -- "What does the C++ community say?"\n'
        "    |                    Queries Pinecone + SQLite.\n"
        "    |                    Returns: structured context (no code).\n"
        "    v\n"
        '[2] CODING AGENT     -- "Implement the fix based on context."\n'
        "    |                    Writes code, creates clean diffs.\n"
        "    |                    Returns: changed files + commit message.\n"
        "    v\n"
        '[3] REVIEW AGENT     -- "Is this code correct and complete?"\n'
        "    |                    Runs tests, validates quality.\n"
        '    |                    Returns: "ALL CLEAR" or feedback.\n'
        "    |\n"
        "    +--> ALL CLEAR? --> Create PR on GitHub\n"
        "    +--> FEEDBACK?  --> Back to Coding Agent (max 3 loops)"
    )
    add_code_block(doc, pipeline)

    h_analogy = add_heading_styled(doc, "Professional Specialization Analogies", 3)
    set_keep_with_next(h_analogy)
    set_keep_together(h_analogy)
    b1 = add_bullet(doc, " is like a research librarian -- finds information but never writes the paper", "Knowledge Agent")
    set_keep_with_next(b1)
    b2 = add_bullet(doc, " is like a senior developer -- writes code but doesn't review its own work", "Coding Agent")
    set_keep_with_next(b2)
    b3 = add_bullet(doc, " is like a QA engineer -- validates but never implements", "Review Agent")
    set_keep_with_next(b3)
    add_bullet(doc, " is like a project manager -- coordinates but never codes", "Orchestrator")

    # pageBreakBefore on heading — never creates blank pages
    h_upgrade = add_heading_styled(doc, "Upgrade Path", 3)
    set_page_break_before(h_upgrade)
    add_bullet(doc, " -- responds to check_run failures, fixes build errors", "CI-Fixer Agent")
    add_bullet(doc, " -- generates/updates docs when code changes", "Documentation Agent")
    add_bullet(doc, " -- scans changes for vulnerabilities before PR creation", "Security Agent")

    add_body(doc, "The architecture scales horizontally by adding agents, not by making existing agents more complex.")


def build_correction(doc):
    h = add_heading_styled(doc, "Appendix: Correction Notice", 1)
    set_page_break_before(h)
    add_accent_bar(doc)

    add_body(doc, "Previous incorrect statement (verbal, not in documents):", bold=True)
    add_body(doc, '"Cursor CLI re-indexes the codebase on every run, causing 2-5 minute cold starts"', italic=True, color=RED_ACCENT)

    add_body(doc, "Corrected statement:", bold=True, color=GREEN_ACCENT)
    add_body(doc, (
        "Cursor CLI does not perform codebase indexing. Indexing is exclusively an IDE feature "
        "(docs.cursor.com/context/codebase-indexing). Both Cursor CLI and Claude Code CLI work "
        "file-based without an index. The 2-5 minute delay in the C++ Alliance's current setup "
        "is caused by GitHub Runner VM allocation and npm install, not by any CLI indexing behavior."
    ))

    add_body(doc, "The decision for Claude Code over Cursor CLI is based on:", bold=True)
    add_bullet(doc, "Native Python SDK (claude-agent-sdk) for programmatic FastAPI integration")
    add_bullet(doc, "Built-in subagents and agent teams for multi-agent orchestration")
    add_bullet(doc, "15 hook events functional in headless/CI mode (vs 8 IDE-only events in Cursor)")
    add_body(doc, "Not on indexing performance.", bold=True)


def build_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SG Consulting  |  Architecture Decisions Q&A  |  February 2026  |  INTERNAL")
        run.font.size = Pt(7)
        run.font.color.rgb = WARM_GRAY
        run.font.name = "Calibri"


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = CHARCOAL

    build_cover(doc)
    build_q1(doc)
    build_q2(doc)
    build_q3(doc)
    build_q4(doc)
    build_q5(doc)
    build_q6(doc)
    build_correction(doc)
    build_footer(doc)

    out = "QA-ARCHITECTURE-DECISIONS.docx"
    doc.save(out)
    print(f"[OK] Generated: {out}")
    print(f"     Pages: ~12+ (estimated)")
    print(f"     Questions: 6 + Correction Notice")
    print(f"     Tables: 11 styled tables")


if __name__ == "__main__":
    main()
